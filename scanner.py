import os
import re
import time
import uuid

from datetime import datetime

import importlib.util
import tempfile

HAS_WORD_COM = (
    importlib.util.find_spec("pythoncom") is not None
    and importlib.util.find_spec("win32com") is not None
)

if HAS_WORD_COM:
    import pythoncom
    import win32com.client
else:
    pythoncom = None
    win32com = None

from docx import Document

from models import Protocol, db


SCAN_PATH = r"\\Admin\рабочая"

VALID_EXTENSIONS = [".docx", ".doc"]

IGNORE_WORDS = [
    "перечень протоколов",
    "реестр протоколов",
    "журнал протоколов"
]

MONTHS = {
    "января": "01",
    "февраля": "02",
    "марта": "03",
    "апреля": "04",
    "мая": "05",
    "июня": "06",
    "июля": "07",
    "августа": "08",
    "сентября": "09",
    "октября": "10",
    "ноября": "11",
    "декабря": "12",
}

def is_protocol(filename):

    name = filename.lower()

    # временные файлы Word
    if name.startswith("~$"):
        return False

    if "протокол" not in name:
        return False

    ignore_words = [
        "перечень протоколов",
        "реестр протоколов",
        "журнал протоколов"
    ]

    for word in ignore_words:

        if word in name:
            return False

    return True


def read_word_file(path):
    try:

        # DOCX
        if path.lower().endswith(".docx"):
            return extract_text_from_docx(path)

        # DOC
        elif path.lower().endswith(".doc"):
            temp_docx = convert_doc_to_docx(path)

            if not temp_docx:
                return ""

            try:
                return extract_text_from_docx(temp_docx)
            finally:
                try:
                    os.remove(temp_docx)
                except Exception:
                    pass

    except Exception as e:

        print(f"Ошибка чтения файла {path}: {e}")

        return ""

    return ""


def extract_text_from_docx(path):
    doc = Document(path)
    text = [p.text for p in doc.paragraphs]

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text.append(cell.text)

    return clean_text("\n".join(text))


def clean_text(text):
    text = text.replace("\x07", " ")
    # удаляем мусор, но сохраняем \n и \t
    text = re.sub(r"[\x00-\x08\x0b-\x1f]", " ", text)
    # убираем лишние пробелы
    text = re.sub(r"[ \t]+", " ", text)
    # чистим пустые строки
    text = re.sub(r"\n+", "\n", text)
    return text.strip()


def convert_doc_to_docx(path):

    if not HAS_WORD_COM:
        print("Конвертация DOC недоступна: не установлен pywin32/Word COM")
        return None

    if len(path) > 240:
        print(f"Слишком длинный путь: {path}")
        return None

    word = None
    doc = None

    try:

        pythoncom.CoInitialize()

        word = win32com.client.DispatchEx("Word.Application")

        word.Visible = False
        word.DisplayAlerts = 0
        word.AutomationSecurity = 3

        doc = word.Documents.Open(
            path,
            ReadOnly=True,
            ConfirmConversions=False,
            AddToRecentFiles=False,
            Visible=False,
            OpenAndRepair=True,
            NoEncodingDialog=True
        )

        temp_docx = os.path.join(
            tempfile.gettempdir(),
            f"{uuid.uuid4()}.docx"
        )

        doc.SaveAs(
            temp_docx,
            FileFormat=16
        )

        return temp_docx

    except Exception as e:

        print(f"Ошибка конвертации DOC: {path}: {e}")

        return None

    finally:

        try:
            if doc:
                doc.Close(False)
        except Exception:
            pass

        try:
            if word:
                word.Quit()
        except Exception:
            pass

        pythoncom.CoUninitialize()


def extract_date(text):

    patterns = [

        # 17.07.2024
        r"(\d{1,2}\.\d{1,2}\.\s*\d{4})",

        # 2008-12-15
        r"(\d{4}-\d{2}-\d{2})",

        # 24 июня 2010
        r"(\d{1,2})\s+([А-Яа-яЁё]+)\s+(\d{4})",
    ]

    # -----------------------------
    # dd.mm.yyyy
    # -----------------------------
    match = re.search(
        patterns[0],
        text,
        re.IGNORECASE
    )

    if match:

        try:

            date_str = re.sub(
                r"\s+",
                "",
                match.group(1)
            )

            return datetime.strptime(
                date_str,
                "%d.%m.%Y"
            ).date()

        except Exception as e:

            print("Ошибка dd.mm.yyyy:", e)

    # -----------------------------
    # yyyy-mm-dd
    # -----------------------------
    match = re.search(
        patterns[1],
        text,
        re.IGNORECASE
    )

    if match:

        try:

            return datetime.strptime(
                match.group(1),
                "%Y-%m-%d"
            ).date()

        except Exception as e:

            print("Ошибка yyyy-mm-dd:", e)

    # -----------------------------
    # 24 июня 2010
    # -----------------------------
    match = re.search(
        patterns[2],
        text,
        re.IGNORECASE
    )

    if match:

        try:

            day = match.group(1).zfill(2)

            month_name = match.group(2).lower()

            year = match.group(3)

            month = MONTHS.get(month_name)

            if month:

                return datetime.strptime(
                    f"{day}.{month}.{year}",
                    "%d.%m.%Y"
                ).date()

        except Exception as e:

            print("Ошибка текстовой даты:", e)

    return None
    
def extract_protocol_data(path):

    print(f"\nИзвлечение данных: {path}")
    text = read_word_file(path)

    if not text:
        print("Текст документа пуст")
        return None

    print("Документ прочитан")
    filename = os.path.basename(path)

    protocol_name = filename.replace(".docx", "").replace(".doc", "")

    object_match = re.search(
        r"Объект\s*[:\-]?\s*(.*?)(?:Дата проведения|Протокол №|Испытания|Назначение|Проект|Основные характеристики)",
        text,
        re.IGNORECASE | re.DOTALL
    )
    print("Поиск даты...")
        
    test_date = extract_date(text)
        
    protocol_match = re.search(
        r"ПРОТОКОЛ\s*(?:№|N|No)?\s*([A-Za-zА-Яа-я0-9\-\/]+)",
        text,
        re.IGNORECASE
    )
    
    if not protocol_match:
        print("Номер протокола НЕ найден")
    else:
        print(f"Номер найден: {protocol_match.group(1)}")
        
    engineers = []

    try:

        # участок между "произвели" и "Руководитель"
        engineers_block = re.search(
            r"произвел(.*?)уководит",
            text,
            re.IGNORECASE | re.DOTALL
        )

        if engineers_block:

            block = engineers_block.group(1)

            found_engineers = re.findall(
                r"/\s*([А-ЯЁ][а-яё\-]+?\s+[А-Я]\.[А-Я]\.)\s*/",
                block
            )

            engineers = sorted(set(found_engineers))

    except Exception as e:

        print("Ошибка поиска инженеров:", e)

    object_name = re.sub(r"\s+", " ", object_match.group(1)).strip() if object_match else ""

    protocol_number = ""
    if protocol_match:
        protocol_number = protocol_match.group(1).strip()
        protocol_number = re.sub(r"\s+", " ", protocol_number)
        # защита от огромных значений
        if len(protocol_number) > 50:
            protocol_number = protocol_number[:50]

    try:
             
        print(f"Дата: {test_date}")

    except:
        return None

    if test_date.year < 2000:
        return None

    engineers = ", ".join(engineers)

    return {
        "protocol_number": protocol_number,
        "protocol_name": protocol_name,
        "object_name": object_name,
        "test_date": test_date,
        "engineers": engineers,
        "file_path": path
    }


def save_protocol(data):

    existing = Protocol.query.filter_by(
        file_path=data["file_path"]
    ).first()

    if existing:
        existing.protocol_number = data["protocol_number"]
        existing.protocol_name = data["protocol_name"]
        existing.object_name = data["object_name"]
        existing.test_date = data["test_date"]
        existing.engineers = data["engineers"]
        existing.last_modified = data["last_modified"]
        db.session.commit()
        return

    protocol = Protocol(**data)

    db.session.add(protocol)
    db.session.commit()

    print(f"Добавлен: {data['protocol_name']}")


def scan_files(app):

    with app.app_context():

        print("Сканирование файлов...")
        scanned_paths = set()

        for root, dirs, files in os.walk(SCAN_PATH):

            for file in files:

                try:

                    filename_lower = file.lower()

                    if file.startswith("~$"):
                        continue

                    if not is_protocol(file):
                        continue

                    if not (
                        filename_lower.endswith(".docx")
                        or filename_lower.endswith(".doc")
                    ):
                        print("Пропуск: неверное расширение")
                        continue

                    ext = os.path.splitext(file)[1].lower()

                    if ext not in VALID_EXTENSIONS:
                        print("Пропуск: расширение не разрешено")
                        continue

                    path = os.path.join(root, file)
                    print(f"Путь: {path}")
                    scanned_paths.add(path)

                    if len(path) > 240:
                        print("Пропуск: слишком длинный путь")
                        continue

                    last_modified = datetime.utcfromtimestamp(
                        os.path.getmtime(path)
                    )

                    existing = Protocol.query.filter_by(
                        file_path=path
                    ).first()

                    if (
                        existing
                        and existing.last_modified
                        and existing.last_modified >= last_modified
                    ):
                        print("Пропуск: файл не изменился")
                        continue

                    print("Чтение документа...")
                    data = extract_protocol_data(path)

                    if data: 
                        print("Успешно извлечены данные:")
                        data["last_modified"] = last_modified
                        save_protocol(data)
                    else:
                        print("Не удалось извлечь данные")

                except Exception as e:
                    print(f"Ошибка обработки файла {file}:")
                    print(e)

        try:
            db_paths = {
                row[0]
                for row in db.session.query(Protocol.file_path).all()
                if row[0]
            }

            deleted_paths = db_paths - scanned_paths

            if deleted_paths:
                deleted_count = Protocol.query.filter(
                    Protocol.file_path.in_(deleted_paths)
                ).delete(synchronize_session=False)
                db.session.commit()
                print(f"Удалено из БД (файлы отсутствуют): {deleted_count}")

        except Exception as e:
            db.session.rollback()
            print("Ошибка удаления отсутствующих файлов:", e)

        print("Сканирование завершено")


def background_scan(app):

    while True:

        try:
            scan_files(app)

        except Exception as e:
            print("Ошибка сканирования:", e)

        time.sleep(300)
